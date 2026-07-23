"""
MaxKB 核心模块源码注释 - 教学版
按源码结构组织，配套 PPT 中的讲解
"""

# ============================================================================
# 模块 1：应用入口（apps/application/views.py · 模型视图）
# 来源：MaxKB 源码 apps/application/views.py
# ============================================================================

APPLICATION_VIEW_CODE = '''
# === MaxKB 应用层 - API 视图简化版 ===
# 位置：apps/application/views.py
# 完整路径：https://github.com/1Panel-dev/MaxKB/blob/main/apps/application/views.py

from rest_framework.views import APIView
from rest_framework.response import Response
from django.contrib.auth.models import User
from apps.application.models import Application
from apps.application.serializers import ApplicationSerializer
from apps.application.chat_service import chat_service

class ApplicationView(APIView):
    """应用管理视图 - 提供 CRUD + 聊天 API"""

    def post(self, request, application_id=None):
        """发起对话 - 核心聊天接口"""
        # Step 1: 校验应用存在
        app = Application.objects.get(id=application_id)

        # Step 2: 获取用户问题
        question = request.data.get('question')
        if not question:
            return Response({'error': '问题不能为空'}, status=400)

        # Step 3: 调用聊天服务（RAG + Workflow）
        # 内部走 RAG Pipeline / Workflow Engine
        answer = chat_service.chat(
            application=app,
            question=question,
            history=request.data.get('history', []),
            user_id=request.user.id if request.user.is_authenticated else None,
        )

        # Step 4: 序列化返回（含参考文档）
        return Response({
            'answer': answer.content,
            'reference_documents': [
                {'name': d.name, 'content': d.content[:500]}
                for d in answer.references
            ],
            'chat_id': answer.chat_id,
        })
'''


# ============================================================================
# 模块 2：RAG 核心 Pipeline（apps/knowledge/chat_pipeline.py）
# ============================================================================

RAG_PIPELINE_CODE = '''
# === RAG Pipeline 完整实现简化版 ===
# 位置：apps/knowledge/chat_pipeline.py

import logging
from typing import List
from langchain.schema import Document
from apps.embedding.models import Embedding
from apps.knowledge.models import KnowledgeBase, Document as KbDocument
from apps.models.models import Model
from langchain.vectorstores import PGVector
from langchain.embeddings import OpenAIEmbeddings

logger = logging.getLogger(__name__)

class RAGPipeline:
    """RAG 检索增强生成管道 - MaxKB 核心"""

    def __init__(self, knowledge_base: KnowledgeBase, model: Model,
                 embedding_model: Embedding):
        self.kb = knowledge_base
        self.llm = self._build_llm(model)
        self.embedding = self._build_embedding(embedding_model)
        self.vector_store = self._build_vector_store()

    def _build_llm(self, model: Model):
        """构建 LLM，根据 model.provider 选择不同实现"""
        # DeepSeek / OpenAI / Claude / 通义千问 ...
        # 统一封装为 ChatOpenAI compatible 接口
        from langchain.chat_models import ChatOpenAI
        return ChatOpenAI(
            model=model.name,
            openai_api_key=model.api_key,
            openai_api_base=model.base_url,
            temperature=0.7,
        )

    def _build_embedding(self, emb: Embedding):
        """构建 Embedding 模型"""
        return OpenAIEmbeddings(
            model=emb.model_name,
            openai_api_key=emb.api_key,
        )

    def _build_vector_store(self):
        """构建 PGVector 向量库"""
        return PGVector(
            connection_string=self.kb.pg_conn_str,
            embedding_function=self.embedding,
            collection_name=f'kb_{self.kb.id}',
        )

    def query(self, question: str, top_k: int = 5) -> dict:
        """核心查询方法：召回 → 组装 → 生成"""
        # Step 1: 向量检索 Top-K
        docs_with_score = self.vector_store.similarity_search_with_score(
            question, k=top_k
        )

        # Step 2: 构造 Prompt
        context = self._build_context(docs_with_score)
        prompt = self._build_prompt(context, question)

        # Step 3: LLM 生成
        from langchain.schema import HumanMessage
        answer = self.llm([HumanMessage(content=prompt)])

        return {
            'answer': answer.content,
            'references': [
                {'content': d.page_content, 'score': score,
                 'metadata': d.metadata}
                for d, score in docs_with_score
            ],
        }
'''


# ============================================================================
# 模块 3：文档处理（apps/knowledge/handle/）
# ============================================================================

DOCUMENT_HANDLER_CODE = '''
# === 文档处理管道 - PDF/Word/MD 解析 + 拆分 + 向量化 ===
# 位置：apps/knowledge/handle/document_handler.py

from typing import List
import re
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter

class DocumentHandler:
    """文档处理核心"""

    SUPPORTED_EXTENSIONS = ['.pdf', '.docx', '.md', '.txt', '.csv', '.xlsx']

    def __init__(self, chunk_size: int = 500, chunk_overlap: int = 50):
        self.splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\\n\\n", "\\n", "。", "！", "？", " ", ""],
        )

    def parse(self, file_path: str) -> str:
        """根据扩展名解析文档"""
        ext = '.' + file_path.rsplit('.', 1)[-1].lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f'不支持的文件类型: {ext}')

        if ext == '.pdf':
            return self._parse_pdf(file_path)
        elif ext == '.docx':
            return self._parse_docx(file_path)
        elif ext == '.md':
            return self._parse_md(file_path)
        elif ext == '.txt':
            return self._parse_txt(file_path)
        elif ext in ['.csv', '.xlsx']:
            return self._parse_table(file_path)

    def _parse_pdf(self, file_path: str) -> str:
        """PDF 解析 - 使用 pdfplumber"""
        import pdfplumber
        text = ''
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text += page.extract_text() or ''
                text += '\\n\\n'
        return text

    def _parse_docx(self, file_path: str) -> str:
        """Word 解析"""
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)
        return '\\n'.join(p.text for p in doc.paragraphs)

    def split(self, text: str) -> List[Document]:
        """文本拆分 - 按段落 + 标点"""
        chunks = self.splitter.split_text(text)
        return [Document(page_content=c, metadata={'chunk_index': i})
                for i, c in enumerate(chunks)]

    def process(self, file_path: str) -> List[Document]:
        """完整处理流程：解析 → 拆分"""
        text = self.parse(file_path)
        return self.split(text)
'''


# ============================================================================
# 模块 4：工作流引擎（apps/workflow/serializers/）
# ============================================================================

WORKFLOW_ENGINE_CODE = '''
# === 工作流节点定义 ===
# 位置：apps/workflow/serializers/node/

from dataclasses import dataclass, field
from typing import Dict, Any, List, Optional

@dataclass
class WorkflowNode:
    """工作流节点基类"""
    id: str
    type: str                          # start / llm / knowledge_retrieval / http / code
    name: str
    config: Dict[str, Any] = field(default_factory=dict)
    next_nodes: List[str] = field(default_factory=list)

class StartNode(WorkflowNode):
    """开始节点"""
    type: str = "start"

class LLMNode(WorkflowNode):
    """LLM 对话节点"""
    type: str = "llm"

class KnowledgeRetrievalNode(WorkflowNode):
    """知识库检索节点"""
    type: str = "knowledge_retrieval"

class HTTPNode(WorkflowNode):
    """HTTP 请求节点"""
    type: str = "http"

class CodeNode(WorkflowNode):
    """自定义代码节点"""
    type: str = "code"

class ConditionNode(WorkflowNode):
    """条件分支节点"""
    type: str = "condition"

# === 工作流执行器 ===
# 位置：apps/workflow/engine.py

class WorkflowEngine:
    """工作流执行引擎"""

    def execute(self, workflow_json: dict, inputs: dict) -> dict:
        """根据工作流 JSON 执行"""
        nodes = {n['id']: self._build_node(n)
                 for n in workflow_json['nodes']}
        edges = workflow_json['edges']

        # 从 start 节点开始
        current = 'start'
        context = {'inputs': inputs, 'outputs': {}}

        while current:
            node = nodes[current]
            result = self._run_node(node, context)

            # 保存到 context
            context['outputs'][current] = result

            # 找到下一个节点
            next_nodes = self._find_next(nodes, edges, current, result)

            if not next_nodes:
                break
            current = next_nodes[0]

        return context['outputs']

    def _build_node(self, node_def: dict) -> WorkflowNode:
        """根据 type 创建对应节点实例"""
        cls_map = {
            'start': StartNode,
            'llm': LLMNode,
            'knowledge_retrieval': KnowledgeRetrievalNode,
            'http': HTTPNode,
            'code': CodeNode,
            'condition': ConditionNode,
        }
        cls = cls_map[node_def['type']]
        return cls(**node_def)

    def _run_node(self, node: WorkflowNode, context: dict) -> Any:
        """执行单个节点"""
        if isinstance(node, LLMNode):
            return self._run_llm(node, context)
        elif isinstance(node, KnowledgeRetrievalNode):
            return self._run_retrieval(node, context)
        # ... 其他节点类型
        return None
'''


# ============================================================================
# 模块 5：MCP 客户端（apps/tools/mcp/）
# ============================================================================

MCP_CLIENT_CODE = '''
# === MCP 客户端（Model Context Protocol）===
# 位置：apps/tools/mcp/client.py

import json
import requests
from typing import Dict, Any, List
from pydantic import BaseModel

class MCPTool(BaseModel):
    """MCP 工具定义"""
    name: str
    description: str
    parameters: Dict[str, Any]

class MCPClient:
    """MCP 协议客户端"""

    def __init__(self, server_url: str, api_key: str):
        self.server_url = server_url
        self.api_key = api_key
        self.tools: List[MCPTool] = []

    def discover_tools(self) -> List[MCPTool]:
        """发现 Server 提供的所有工具"""
        response = requests.get(
            f'{self.server_url}/mcp/tools',
            headers={'Authorization': f'Bearer {self.api_key}'}
        )
        tools_data = response.json()['tools']
        self.tools = [MCPTool(**t) for t in tools_data]
        return self.tools

    def call_tool(self, name: str, params: Dict[str, Any]) -> Dict[str, Any]:
        """调用 MCP 工具"""
        response = requests.post(
            f'{self.server_url}/mcp/call',
            headers={
                'Authorization': f'Bearer {self.api_key}',
                'Content-Type': 'application/json',
            },
            json={'name': name, 'params': params},
        )
        return response.json()

    def to_openai_function(self, tool: MCPTool) -> dict:
        """转换为 OpenAI Function Calling 格式"""
        return {
            'name': tool.name,
            'description': tool.description,
            'parameters': tool.parameters,
        }


# === 使用示例 ===
# 在工作流中拖入「MCP 节点」→ 选择工具 → 自动转换
client = MCPClient(server_url='http://mcp.example.com', api_key='sk-xxx')
tools = client.discover_tools()
result = client.call_tool('search_documents', {'query': 'MaxKB'})
'''


# ============================================================================
# 模块 6：模型接入适配层
# ============================================================================

MODEL_ADAPTER_CODE = '''
# === 多模型统一接入（模型中立）===
# 位置：apps/models/models/model_provider.py

from typing import Dict, Any
from langchain.chat_models import ChatOpenAI, ChatAnthropic

class ModelProvider:
    """统一模型提供者 - 模型中立核心"""

    SUPPORTED_PROVIDERS = {
        'deepseek': {
            'class': ChatOpenAI,
            'base_url': 'https://api.deepseek.com/v1',
            'default_model': 'deepseek-chat',
        },
        'openai': {
            'class': ChatOpenAI,
            'base_url': 'https://api.openai.com/v1',
            'default_model': 'gpt-4-turbo',
        },
        'claude': {
            'class': ChatAnthropic,
            'base_url': 'https://api.anthropic.com',
            'default_model': 'claude-3-5-sonnet',
        },
        'qwen': {
            'class': ChatOpenAI,                # 通义千问兼容 OpenAI 协议
            'base_url': 'https://dashscope.aliyuncs.com/compatible-mode/v1',
            'default_model': 'qwen-max',
        },
        'zhipu': {
            'class': ChatOpenAI,
            'base_url': 'https://open.bigmodel.cn/api/paas/v4',
            'default_model': 'glm-4',
        },
    }

    def get_llm(self, provider: str, model_name: str = None,
                api_key: str = None, **kwargs) -> Any:
        """根据 provider 获取 LLM 实例"""
        cfg = self.SUPPORTED_PROVIDERS.get(provider)
        if not cfg:
            raise ValueError(f'不支持的模型: {provider}')

        # 构造 LLM 实例
        model_name = model_name or cfg['default_model']

        if provider == 'claude':
            return cfg['class'](model=model_name, anthropic_api_key=api_key, **kwargs)
        else:
            return cfg['class'](
                model=model_name,
                openai_api_key=api_key,
                openai_api_base=cfg['base_url'],
                **kwargs
            )
'''


# ============================================================================
# 模块 7：Embedding & 向量化
# ============================================================================

EMBEDDING_CODE = '''
# === Embedding & 向量化模块 ===
# 位置：apps/embedding/models/

from langchain.embeddings import OpenAIEmbeddings, HuggingFaceEmbeddings

class EmbeddingFactory:
    """Embedding 工厂 - 支持多种 embedding 模型"""

    def create(self, provider: str, model_name: str, api_key: str = None,
               device: str = 'cpu'):
        """创建 Embedding 实例"""
        if provider == 'openai':
            return OpenAIEmbeddings(
                model=model_name or 'text-embedding-3-small',
                openai_api_key=api_key,
            )
        elif provider == 'huggingface':
            return HuggingFaceEmbeddings(
                model_name=model_name or 'BAAI/bge-large-zh-v1.5',
                model_kwargs={'device': device},
                encode_kwargs={'normalize_embeddings': True},
            )
        elif provider == 'bce':
            # 网易 BCE Embedding
            return OpenAIEmbeddings(
                model=model_name,
                openai_api_key=api_key,
                openai_api_base='https://api.bce.baidu.com/v1',
            )
        # ...

# === PGVector 存储示例 ===
from langchain.vectorstores import PGVector

def create_vector_store(docs, embedding, conn_str):
    """创建 pgvector 向量库"""
    return PGVector.from_documents(
        documents=docs,
        embedding=embedding,
        connection_string=conn_str,
        pre_delete_collection=True,
    )
'''


# ============================================================================
# 模块 8：前端 Vue.js 组件（API 调用）
# ============================================================================

FRONTEND_COMPONENT_CODE = '''
// === 前端 Vue 3 组件 - MaxKB Web 控制台 ===
// 位置：ui/src/views/application/index.vue (简化版)

// <template>
<template>
  <div class="chat-container">
    <div class="messages">
      <MessageBubble v-for="msg in messages" :key="msg.id" :msg="msg" />
    </div>
    <div class="input-area">
      <el-input v-model="question" type="textarea" :rows="3"
        placeholder="请输入您的问题..." />
      <el-button type="primary" @click="send" :loading="loading">
        发送
      </el-button>
    </div>
  </div>
</template>

<script setup>
import { ref } from 'vue'
import { chatApi } from '@/api/chat'
import MessageBubble from './MessageBubble.vue'

const question = ref('')
const messages = ref([])
const loading = ref(false)

const send = async () => {
  if (!question.value.trim()) return
  const userMsg = {
    id: Date.now(),
    role: 'user',
    content: question.value,
  }
  messages.value.push(userMsg)
  loading.value = true
  try {
    const res = await chatApi.send({
      application_id: 'xxx',
      question: question.value,
    })
    messages.value.push({
      id: Date.now() + 1,
      role: 'assistant',
      content: res.answer,
      references: res.reference_documents,
    })
    question.value = ''
  } finally {
    loading.value = false
  }
}
</script>
'''


# ============ 主文件 ============

if __name__ == '__main__':
    print('MaxKB 源码模块注释 · 教学版')
    print('=' * 60)
    modules = [
        ('模块 1', '应用入口 application/views.py', APPLICATION_VIEW_CODE),
        ('模块 2', 'RAG Pipeline knowledge/chat_pipeline.py', RAG_PIPELINE_CODE),
        ('模块 3', '文档处理 knowledge/handle/document_handler.py', DOCUMENT_HANDLER_CODE),
        ('模块 4', '工作流引擎 workflow/engine.py', WORKFLOW_ENGINE_CODE),
        ('模块 5', 'MCP 客户端 tools/mcp/client.py', MCP_CLIENT_CODE),
        ('模块 6', '模型接入 models/model_provider.py', MODEL_ADAPTER_CODE),
        ('模块 7', 'Embedding 模块 embedding/models/', EMBEDDING_CODE),
        ('模块 8', '前端 Vue 组件 ui/src/views/application/', FRONTEND_COMPONENT_CODE),
    ]
    for num, path, _ in modules:
        print(f'{num} | {path}')
    print()
    print(f'共 {len(modules)} 个核心模块')
    print(f'目录：/root/.openclaw/workspace/projects/MaxKB教学案例/code/')
